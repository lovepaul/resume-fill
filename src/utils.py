"""
工具函数 — 字体设置、段落操作等底层工具
"""
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .config import FONT_NAME, FONT_SIZE


def set_run_font(run, bold=None):
    """设置 run 的字体属性：Microsoft YaHei + 9pt + eastAsia"""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def normalize_all_fonts(doc):
    """遍历文档中所有文本 run，统一设置为 Microsoft YaHei"""
    # 段落
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text.strip():
                set_run_font(r, bold=r.bold)
    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.text.strip():
                            set_run_font(r, bold=r.bold)


def find_heading_idx(doc, heading_text):
    """查找包含指定文本的标题段落索引"""
    for i, p in enumerate(doc.paragraphs):
        if heading_text in p.text:
            return i
    return None


def find_next_heading_idx(doc, start_idx):
    """从 start_idx+1 开始找下一个标题段落索引"""
    for i in range(start_idx + 1, len(doc.paragraphs)):
        style = doc.paragraphs[i].style.name if doc.paragraphs[i].style else ""
        if style.startswith('Heading'):
            return i
    return len(doc.paragraphs)


def remove_paragraphs_between(doc, start_idx, end_idx):
    """从 XML 中删除 start_idx+1 到 end_idx-1 之间的段落"""
    for idx in range(end_idx - 1, start_idx, -1):
        p_element = doc.paragraphs[idx]._element
        p_element.getparent().remove(p_element)


def add_paragraph_after(doc, ref_para, text, style_name='Normal', bold=None):
    """在 ref_para 之后插入一个新段落，统一字体"""
    new_para = doc.add_paragraph(text, style=style_name)
    for run in new_para.runs:
        set_run_font(run, bold=bold)
    ref_para._element.addnext(new_para._element)
    return new_para


def clean_heading(doc, heading_text):
    """清除标题中的模板指令文本，只保留标题本身，并统一字体"""
    idx = find_heading_idx(doc, heading_text)
    if idx is not None:
        p = doc.paragraphs[idx]
        for r in p.runs:
            r.text = ""
        if p.runs:
            p.runs[0].text = heading_text
            set_run_font(p.runs[0], bold=None)


def fill_table_cell(table, row, col, text, bold=None):
    """安全填充表格单元格，统一字体"""
    cell = table.cell(row, col)
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    run = cell.paragraphs[0].add_run(text)
    set_run_font(run, bold=bold)
    return cell


def fill_section(doc, section_heading, items, item_style='List Paragraph'):
    """通用段落区填充：删除占位 → 插入新内容"""
    heading_idx = find_heading_idx(doc, section_heading)
    if heading_idx is None:
        print(f"  ⚠ 未找到标题: {section_heading}")
        return None

    next_heading_idx = find_next_heading_idx(doc, heading_idx)
    remove_paragraphs_between(doc, heading_idx, next_heading_idx)

    heading_para = doc.paragraphs[heading_idx]
    last_para = heading_para

    for item in items:
        last_para = add_paragraph_after(doc, last_para, item, item_style)

    return last_para
