"""
模板填充核心 — 所有 fill_* 函数
"""
import docx
from .utils import (
    fill_table_cell, fill_section, add_paragraph_after,
    find_heading_idx, find_next_heading_idx,
    remove_paragraphs_between, clean_heading, normalize_all_fonts
)


def fill_resource_info(doc, data):
    """填充 Resource Information 表格 (Table 0: 3行×2列)"""
    info = data["resource_info"]
    table = doc.tables[0]

    cell_text = "所在城市:\n语言能力:\n出生日期:\n性别:\n可面试时间:"
    value_text = (
        f"{info['city']}\n{info['language']}\n{info['birth']}\n"
        f"{info['gender']}\n{info['interview_time']}"
    )

    fill_table_cell(table, 0, 0, cell_text, bold=None)
    fill_table_cell(table, 0, 1, value_text, bold=True)
    fill_table_cell(table, 1, 0, "目前状态:", bold=None)
    fill_table_cell(table, 1, 1, info['status'], bold=True)
    fill_table_cell(table, 2, 0, "")
    fill_table_cell(table, 2, 1, "")


def fill_education(doc, data):
    """填充 Education 表格 (Table 1: 4行×2列)"""
    table = doc.tables[1]
    edu_list = data["education"]

    for idx, edu in enumerate(edu_list[:2]):
        row_offset = idx * 2
        fill_table_cell(table, row_offset, 0, edu['period'])
        fill_table_cell(table, row_offset, 1, edu['school'])
        fill_table_cell(table, row_offset + 1, 0, "")
        fill_table_cell(table, row_offset + 1, 1, edu['degree'])

    for idx in range(len(edu_list), 2):
        row_offset = idx * 2
        fill_table_cell(table, row_offset, 0, "")
        fill_table_cell(table, row_offset, 1, "")
        fill_table_cell(table, row_offset + 1, 0, "")
        fill_table_cell(table, row_offset + 1, 1, "")


def fill_employment_history(doc, data):
    """填充 Employment History 表格 (Table 2: 头部+数据行)"""
    table = doc.tables[2]
    emp_list = data["employment_history"]

    data_start_row = 1
    template_data_rows = len(table.rows) - data_start_row

    rows_needed = len(emp_list) - template_data_rows
    for _ in range(rows_needed):
        table.add_row()

    for idx, emp in enumerate(emp_list):
        row = data_start_row + idx
        fill_table_cell(table, row, 0, emp['time'])
        fill_table_cell(table, row, 1, emp['employer'])
        fill_table_cell(table, row, 2, emp['role'])

    for idx in range(len(emp_list), template_data_rows):
        row = data_start_row + idx
        fill_table_cell(table, row, 0, "")
        fill_table_cell(table, row, 1, "")
        fill_table_cell(table, row, 2, "")


def fill_summary(doc, data):
    """填充 Summary 区（List Paragraph 样式）"""
    print("  填充 Summary ...")
    fill_section(doc, "Summary", data["summary_items"], "List Paragraph")


def fill_roles(doc, data):
    """填充 Role and Accomplishment 区"""
    print("  填充 Role and Accomplishment ...")
    heading_idx = find_heading_idx(doc, "Role and Accomplishment")
    if heading_idx is None:
        print("  ⚠ 未找到 Role and Accomplishment 标题")
        return

    clean_heading(doc, "Role and Accomplishment")
    next_heading_idx = find_next_heading_idx(doc, heading_idx)
    remove_paragraphs_between(doc, heading_idx, next_heading_idx)

    heading_para = doc.paragraphs[heading_idx]
    last_para = heading_para

    for role in data["roles"]:
        header_line = f"{role['period']}\t{role['company']}\t{role['title']}"
        last_para = add_paragraph_after(doc, last_para, header_line, "Normal", bold=True)

        for i, resp in enumerate(role["responsibilities"], 1):
            last_para = add_paragraph_after(doc, last_para, f"{i}. {resp}", "Normal", bold=True)

        if role.get("achievements"):
            last_para = add_paragraph_after(doc, last_para, "业绩：", "Normal", bold=True)
            for i, ach in enumerate(role["achievements"], 1):
                last_para = add_paragraph_after(doc, last_para, f"{i}. {ach}", "Normal", bold=True)

        last_para = add_paragraph_after(doc, last_para, "", "Normal")


def fill_projects(doc, data):
    """填充 Project Experience 区"""
    print("  填充 Project Experience ...")
    heading_idx = find_heading_idx(doc, "Project Experience")
    if heading_idx is None:
        print("  ⚠ 未找到 Project Experience 标题")
        return

    clean_heading(doc, "Project Experience")
    remove_paragraphs_between(doc, heading_idx, len(doc.paragraphs))

    heading_para = doc.paragraphs[heading_idx]
    last_para = heading_para

    for proj in data["projects"]:
        last_para = add_paragraph_after(
            doc, last_para, f"{proj['period']}\t{proj['name']}", "Normal"
        )
        last_para = add_paragraph_after(
            doc, last_para, f"项目描述：{proj['description']}", "List Paragraph"
        )
        last_para = add_paragraph_after(
            doc, last_para, f"技术栈：{proj['tech_stack']}", "List Paragraph"
        )
        last_para = add_paragraph_after(doc, last_para, "责任描述：", "List Paragraph")

        for resp in proj["responsibilities"]:
            last_para = add_paragraph_after(doc, last_para, resp, "List Paragraph")

        if proj.get("achievements"):
            for ach in proj["achievements"]:
                last_para = add_paragraph_after(
                    doc, last_para, f"业绩：{ach}", "Normal"
                )

        last_para = add_paragraph_after(doc, last_para, "", "Normal")


def fill_template(template_path, data, output_path):
    """主填充流程"""
    print(f"\n📄 加载模板: {template_path}")
    doc = docx.Document(str(template_path))

    # 清理指令文本
    clean_heading(doc, "Education and Professional Development")
    clean_heading(doc, "Employment History")

    fill_resource_info(doc, data)
    fill_education(doc, data)
    fill_employment_history(doc, data)
    fill_summary(doc, data)
    fill_roles(doc, data)
    fill_projects(doc, data)

    # 全局字体统一
    normalize_all_fonts(doc)

    print(f"\n💾 保存到: {output_path}")
    output_path = str(output_path)
    doc.save(output_path)
    print("✅ 完成！")
